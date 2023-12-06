import streamlit as st
from deep_translator import GoogleTranslator
from forms import google_translate_language_options
import shutil
import os
import base64
import docx
from pdf2docx.main import parse
from docx2pdf import convert
import sys
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter


# 外部ファイルとの通信用ファイル
input_pdf_path = "../media/pdf/input.pdf"
input_word_path = "../media/word/input.docx"
output_pdf_path = "../media/pdf/output.pdf"
output_word_path = "../media/word/output.docx"


def create_or_remove_folder(path):
    if not os.path.exists(path):
        os.makedirs(path)  # フォルダを作成
        print(f"フォルダを作成 : {path}")
    else:
        shutil.rmtree(path)  # フォルダを削除
        os.makedirs(path)  # フォルダを作成
        print(f"フォルダを作成 : {path}")


def save_text_to_file(target_text, filename):
    with open(filename, "w") as file:
        file.write(target_text)


def divide_by_period(text, divide):
    """
    指定文字数以下で句点により区切る関数
    「text」を「divide」字数以下で最大字数になるよう句点で区切る
    """

    def count_characters(text):
        return len(text)

    def back_count_period(target_text):
        """最後尾と最も近い句点で区切る関数"""
        # 文字列を逆順にする
        reversed_text = target_text[::-1]
        # 句点「。」が後ろから何番目にあるかを取得する
        index = reversed_text.find("。")
        # 反転したインデックスを元の文字列のインデックスに変換する
        original_index = len(target_text) - index
        if index == -1:
            original_index -= 1
        return original_index

    def split_japanese_text(text, divide):
        text_count = count_characters(text)
        # st.write(f"text_count : {text_count}")
        if text_count >= divide:
            # 最初の3000文字を抽出
            text_before_3000 = text[:divide]
            # 3000文字目以前で最近傍の句点で区切った文字列の文字数
            before_3000_period_count = back_count_period(text_before_3000)
            # 3000文字目以前で最近傍の句点で区切った文字列を抽出
            text_before_3000_period = text[:before_3000_period_count]
            text_after_3000_period = text[before_3000_period_count:]
            return text_before_3000_period, text_after_3000_period
        else:
            return text, ""

    divide_period_texts = []
    character_count = count_characters(text)
    while character_count > divide:
        before_divide_period, after_divide_period = split_japanese_text(text, divide)
        divide_period_texts.append(before_divide_period)
        text = after_divide_period
        character_count = count_characters(text)
    before_divide_period, after_divide_period = split_japanese_text(text, divide)
    divide_period_texts.append(before_divide_period)
    return divide_period_texts


def main(page_subtitle="<h1>翻訳サイト</h1><p></p>", upload_way="text"):
    # タイトル
    st.write(page_subtitle, unsafe_allow_html=True)

    if upload_way == "text":
        # 翻訳対象の文章を入力
        target_text = st.text_area("翻訳したい文章を入力してください", value="", height=400)

        # 言語の選択フォーム
        input_language = st.selectbox(
            "翻訳元の言語を選択してください", google_translate_language_options[:, 0].tolist(), index=0
        )
        output_language = st.selectbox(
            "翻訳後の言語を選択してください", google_translate_language_options[:, 0].tolist(), index=1
        )

        if target_text != "":
            input_parameter = google_translate_language_options[
                google_translate_language_options[:, 0] == input_language
            ][0, 1]
            output_parameter = google_translate_language_options[
                google_translate_language_options[:, 0] == output_language
            ][0, 1]

            output_text = GoogleTranslator(
                source=input_parameter, target=output_parameter
            ).translate(target_text)

            st.write("")
            st.write("<h5>翻訳結果</h5>", unsafe_allow_html=True)

            # 翻訳結果を表示
            output_text = output_text.replace("\n", "<br>")
            styled_text = f'<div style="background-color: #f0f2f6; padding: 20px; border-radius: 5px; cursor: text; " contenteditable="true">{output_text}</div>'
            st.markdown(styled_text, unsafe_allow_html=True)

    if upload_way == "word" or upload_way == "pdf":
        input_words_list = []
        if upload_way == "pdf":
            uploaded_files = st.file_uploader(
                "PDFをアップロードしてください", type=["pdf"], accept_multiple_files=True
            )
            print(f"uploaded_files : {len(uploaded_files)}")
            if uploaded_files is not None:
                count = 0
                uploaded_file_names = []
                for uploaded_file in uploaded_files:
                    uploaded_file_names.append(uploaded_file.name)
                    count = count + 1
                    input_loop_pdf_path = input_pdf_path.replace(
                        "input.", f"input_{str(count)}."
                    )
                    input_loop_word_path = input_pdf_path.replace(
                        "input.", f"input_{str(count)}."
                    )
                    with open(input_loop_pdf_path, "wb") as pdf_file:
                        pdf_file.write(uploaded_file.read())
                    parse(input_loop_pdf_path, input_loop_word_path)
                    input_word = docx.Document(input_loop_word_path)
                    input_words_list.append(input_word)
        else:
            uploaded_files = st.file_uploader(
                "Wordファイルをアップロードしてください",
                type=["doc", "docx"],
                accept_multiple_files=True,
            )
            if uploaded_files is not None:
                for uploaded_file in uploaded_files:
                    st.write("")
                    input_word = docx.Document(uploaded_file)
                    input_words_list.append(input_word)

        if uploaded_files is not None:
            st.write("")
            count = 0
            for input_word in input_words_list:
                output_word = docx.Document()
                output_text_list = []
                for index, input_paragraph in enumerate(input_word.paragraphs):
                    paragraph_text = input_paragraph.text
                    translated_text = GoogleTranslator(
                        source="auto", target="ja"
                    ).translate(paragraph_text)
                    output_text_list.append(translated_text + "\n")

                    # 新規Wordファイルにも追加
                    output_paragraph = output_word.add_paragraph(translated_text)
                    output_paragraph.alignment = input_paragraph.alignment
                    font_size = input_paragraph.style.font.size
                    if font_size is not None:
                        if len(output_paragraph.runs) > 0:
                            output_paragraph.runs[0].font.size = font_size

                # ループごとに出力ファイル名を定義
                count = count + 1
                output_loop_word_path = output_word_path.replace(
                    "output.", f"output_{str(count)}."
                )
                # Wordファイルを保存
                output_word.save(output_loop_word_path)

                with open(output_loop_word_path, "rb") as file:
                    # st.download_button(
                    #     f"{count}枚目のwordをダウンロード", file.read(), output_loop_word_path
                    # )
                    # Wordファイルをバイト列に変換
                    word_bytes = None
                    with open(output_loop_word_path, "rb") as file:
                        word_bytes = file.read()
                    # Base64でエンコード
                    b64_word = base64.b64encode(word_bytes).decode()
                    uploaded_file_name = os.path.splitext(
                        uploaded_file_names[count - 1]
                    )[0]
                    href_word = f'<a href="data:application/octet-stream;base64,{b64_word}" download="{uploaded_file_name}.docx">{uploaded_file_name}.pdf の翻訳結果をダウンロード</a>'
                    st.markdown(
                        f"{href_word}",
                        unsafe_allow_html=True,
                    )

    if upload_way == "pptx":
        # https://qiita.com/Kent-747/items/21d1dee3d486f38c4b3f
        # https://qiita.com/code_440/items/9998d97b480db82ef738
        # ↑ このサイトが分かりやすい
        print()


if __name__ == "__main__":
    main()
